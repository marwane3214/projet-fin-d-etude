# -*- coding: utf-8 -*-
"""
Generateur de Rapport PFE — Portail CIMR
Style conforme au modele EMSI (exemple Anas Miftah El Idrissi)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
import datetime

# ─────────────────────────────────────────────
# COULEURS
# ─────────────────────────────────────────────
NAVY    = RGBColor(0x1F, 0x38, 0x64)   # Bleu titre principal (hex: 1F3864)
DARK    = RGBColor(0x00, 0x00, 0x00)     # Noir texte corps
GRAY    = RGBColor(0x40, 0x40, 0x40)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
RED_HDR = RGBColor(0xC0, 0x00, 0x00)    # rouge titres chapitre (optionnel)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='AAAAAA', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def set_paragraph_border_bottom(para, color='1F3864', sz='12'):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)

def page_break(doc):
    doc.add_page_break()

def add_run(para, text, bold=False, italic=False, size=12,
            color=DARK, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Times New Roman'
    return r

def body_para(doc, text='', bold=False, italic=False, size=12,
              color=DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before=3, space_after=6, left_indent=0,
              first_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment     = align
    p.paragraph_format.space_before  = Pt(space_before)
    p.paragraph_format.space_after   = Pt(space_after)
    p.paragraph_format.line_spacing  = Pt(16)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p

def heading1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color
    r.font.name = 'Times New Roman'
    set_paragraph_border_bottom(p, color='1F3864', sz='8')
    return p

def heading2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color
    r.font.name = 'Times New Roman'
    return p

def heading3(doc, text, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = color
    r.font.name = 'Times New Roman'
    return p

def bullet_item(doc, text, level=0, bold_prefix='', size=11):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(15)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True; r1.font.size = Pt(size); r1.font.color.rgb = DARK
        r1.font.name = 'Times New Roman'
    if text:
        r2 = p.add_run(text)
        r2.font.size = Pt(size); r2.font.color.rgb = DARK
        r2.font.name = 'Times New Roman'
    return p

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = DARK
    r.font.name = 'Times New Roman'
    return p

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.italic = True; r.font.size = Pt(10); r.font.color.rgb = DARK
    r.font.name = 'Times New Roman'
    return p

def add_logos_header(doc):
    """Bande de logos en haut de section (comme le modele EMSI)"""
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = t.rows[0].cells
    # Logo CIMR (gauche)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run('CIMR')
    r0.bold = True; r0.font.size = Pt(14); r0.font.color.rgb = NAVY
    r0.font.name = 'Times New Roman'
    # Logo EMSI (droite)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run('EMSI')
    r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = NAVY
    r1.font.name = 'Times New Roman'
    r1b = p1.add_run('\nECOLE MAROCAINE DES SCIENCES DE L\'INGENIEUR')
    r1b.font.size = Pt(7); r1b.font.color.rgb = NAVY
    r1b.font.name = 'Times New Roman'
    c0.width = Cm(6); c1.width = Cm(9)
    for c in (c0, c1):
        c.paragraphs[0].paragraph_format.space_before = Pt(4)
        c.paragraphs[0].paragraph_format.space_after  = Pt(4)
    # Enlever les bordures de la table logos
    for cell in [c0, c1]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement('w:tcBorders')
        for side in ('top','bottom','left','right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tcBdr.append(el)
        tcPr.append(tcBdr)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def chapter_title_page(doc, chapter_num, chapter_title):
    """Page de titre de chapitre style EMSI (grande titre centree)"""
    # Espace en haut
    for _ in range(8):
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after  = Pt(0)
    # Ligne superieure double
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_border_bottom(p, color='000000', sz='12')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    # Titre chapitre
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after  = Pt(12)
    r = p2.add_run(f"{chapter_num} : {chapter_title}")
    r.font.size = Pt(28); r.font.color.rgb = NAVY
    r.font.name = 'Times New Roman'
    # Ligne inferieure double
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_border_bottom(p3, color='000000', sz='12')
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(0)

def add_page_number_footer(doc):
    """Ajoute un pied de page avec numero de page style { N }"""
    section = doc.sections[-1]
    footer  = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ligne au-dessus
    set_paragraph_border_bottom(p, color='000000', sz='6')
    run1 = p.add_run('{ ')
    run1.font.size = Pt(11); run1.font.color.rgb = DARK; run1.font.name = 'Times New Roman'
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run_fld = p.add_run()
    run_fld.font.size = Pt(11); run_fld.font.name = 'Times New Roman'
    run_fld._r.append(fldChar1); run_fld._r.append(instrText); run_fld._r.append(fldChar2)
    run2 = p.add_run(' }')
    run2.font.size = Pt(11); run2.font.color.rgb = DARK; run2.font.name = 'Times New Roman'

def simple_table(doc, headers, rows, col_widths=None,
                 hdr_bg='1F3864', hdr_fg=WHITE):
    """Tableau style professionnel"""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        set_cell_bg(c, hdr_bg)
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(3)
        pp.paragraph_format.space_after  = Pt(3)
        r = pp.add_run(h)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = hdr_fg; r.font.name = 'Times New Roman'
    for ri, row_data in enumerate(rows):
        bg = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        row = tbl.add_row()
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            pp = cell.paragraphs[0]
            pp.paragraph_format.space_before = Pt(3)
            pp.paragraph_format.space_after  = Pt(3)
            r = pp.add_run(str(cell_text))
            r.font.size = Pt(10); r.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def use_case_table(doc, titre, but, acteur, pre, scenario_p, scenario_a):
    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'
    labels = ["Cas d'utilisation :", "But :", "Acteur systeme :",
              "Pre-condition :", "Scenario principal :", "Scenario alternatif :"]
    values = [titre, but, acteur, pre, scenario_p, scenario_a]
    for i in range(6):
        c0, c1 = tbl.rows[i].cells
        set_cell_bg(c0, 'DDEEFF')
        pp0 = c0.paragraphs[0]
        r0 = pp0.add_run(labels[i])
        r0.bold = True; r0.font.size = Pt(10); r0.font.name = 'Times New Roman'
        pp1 = c1.paragraphs[0]
        pp1.paragraph_format.space_before = Pt(2)
        pp1.paragraph_format.space_after  = Pt(2)
        if isinstance(values[i], list):
            for item in values[i]:
                ip = c1.add_paragraph()
                ir = ip.add_run('- ' + item)
                ir.font.size = Pt(10); ir.font.name = 'Times New Roman'
        else:
            r1 = pp1.add_run(values[i])
            r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    # Fusion premiere ligne
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    set_cell_bg(tbl.rows[0].cells[0], '1F3864')
    pp_t = tbl.rows[0].cells[0].paragraphs[0]
    pp_t.clear()
    r_t = pp_t.add_run(titre)
    r_t.bold = True; r_t.font.size = Pt(11)
    r_t.font.color.rgb = WHITE; r_t.font.name = 'Times New Roman'
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

# ════════════════════════════════════════════════════════════════════════
#  DOCUMENT
# ════════════════════════════════════════════════════════════════════════
doc = Document()

# Marges
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# Police par defaut
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

add_page_number_footer(doc)

# ════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ════════════════════════════════════════════════════════════════════════
# En-tete logos (couverture)
t_cover = doc.add_table(rows=1, cols=2)
t_cover.style = 'Table Grid'
cc0, cc1 = t_cover.rows[0].cells
# CIMR gauche
pc0 = cc0.paragraphs[0]
rc0 = pc0.add_run('CIMR\nCaisse Interprofessionnelle\nMarocaine de Retraite')
rc0.bold = True; rc0.font.size = Pt(13); rc0.font.color.rgb = NAVY
rc0.font.name = 'Times New Roman'
pc0.alignment = WD_ALIGN_PARAGRAPH.CENTER
# EMSI droite
pc1 = cc1.paragraphs[0]
pc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc1 = pc1.add_run('EMSI\n')
rc1.bold = True; rc1.font.size = Pt(13); rc1.font.color.rgb = NAVY; rc1.font.name = 'Times New Roman'
rc1b = pc1.add_run('ECOLE MAROCAINE DES\nSCIENCES DE L\'INGENIEUR\nMembre de HONORIS UNITED UNIVERSITIES')
rc1b.font.size = Pt(8); rc1b.font.color.rgb = NAVY; rc1b.font.name = 'Times New Roman'
cc0.width = Cm(7); cc1.width = Cm(8)
for c in (cc0, cc1):
    c.paragraphs[0].paragraph_format.space_before = Pt(8)
    c.paragraphs[0].paragraph_format.space_after  = Pt(8)
# Remove borders cover table
for cell in [cc0, cc1]:
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tcBdr.append(el)
    tcPr.append(tcBdr)

doc.add_paragraph()

# Titre principal
pt = doc.add_paragraph()
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt.paragraph_format.space_before = Pt(10)
rt = pt.add_run('RAPPORT DE PROJET DE FIN D\'ETUDES')
rt.bold = True; rt.underline = True; rt.font.size = Pt(20)
rt.font.color.rgb = DARK; rt.font.name = 'Times New Roman'

ps = doc.add_paragraph()
ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = ps.add_run('4eme Annee en Ingenierie Informatique et Reseaux')
rs.italic = True; rs.font.size = Pt(12); rs.font.color.rgb = DARK
rs.font.name = 'Times New Roman'

# Ligne separatrice
doc.add_paragraph()
pl1 = doc.add_paragraph()
set_paragraph_border_bottom(pl1, '000000', '6')

# Titre du projet (encadre)
doc.add_paragraph()
pt2 = doc.add_paragraph()
pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt2.paragraph_format.space_before = Pt(10)
pt2.paragraph_format.space_after  = Pt(10)
rt2 = pt2.add_run('Developpement d\'un Portail de Gestion\ndes Droits a la Retraite')
rt2.bold = True; rt2.font.size = Pt(18); rt2.font.color.rgb = DARK
rt2.font.name = 'Times New Roman'

pl2 = doc.add_paragraph()
set_paragraph_border_bottom(pl2, '000000', '6')
doc.add_paragraph()

# Realise par
preal = doc.add_paragraph()
preal.paragraph_format.left_indent = Cm(1)
add_run(preal, 'Realise par :', bold=True, italic=True, size=12)
doc.add_paragraph().paragraph_format.left_indent = Cm(1)
pname = doc.add_paragraph()
pname.paragraph_format.left_indent = Cm(4)
add_run(pname, 'Mharrech Iliass', italic=True, size=12)

doc.add_paragraph()

# Tuteurs
ptut = doc.add_paragraph()
ptut.paragraph_format.left_indent = Cm(1)
add_run(ptut, 'Tuteur(s) :', bold=True, italic=True, size=12)

penc1 = doc.add_paragraph()
penc1.paragraph_format.left_indent = Cm(3)
add_run(penc1, 'Encadrant Professionnel : Mr. [Nom Encadrant]', italic=True, size=12)

penc2 = doc.add_paragraph()
penc2.paragraph_format.left_indent = Cm(3)
add_run(penc2, 'Encadrant Pedagogique :   Mr. [Nom Encadrant]', italic=True, size=12)

doc.add_paragraph()

# Organisme
porg = doc.add_paragraph()
porg.paragraph_format.left_indent = Cm(1)
add_run(porg, 'Au sein de : ', bold=True, italic=True, size=12)

porgname = doc.add_paragraph()
porgname.paragraph_format.left_indent = Cm(3)
porgname.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(porgname, 'Caisse Interprofessionnelle Marocaine de Retraite (CIMR)', italic=True, size=12)

doc.add_paragraph()
doc.add_paragraph()

# Annee
pyr = doc.add_paragraph()
pyr.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(pyr, 'Annee universitaire : 2025/2026', bold=True, size=12)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# DEDICACES
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
h_ded = doc.add_heading('Dedicaces', level=1)
for r in h_ded.runs:
    r.font.name = 'Times New Roman'
    r.font.color.rgb = DARK
h_ded.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

deds = [
    "Au nom d'Allah le tout misericordieux, le tres misericordieux",
    "A mes chers parents, pour leur amour inconditionnel, leur soutien infaillible et leurs sacrifices qui ont guide chacun de mes pas.",
    "A ma famille, pour son encouragement constant tout au long de mon parcours academique.",
    "A mes amis et camarades de promotion, pour les moments partages, les discussions stimulantes et l'amitie precieuse.",
    "Enfin, a toutes les personnes qui ont contribue de pres ou de loin au bon deroulement de ce projet, je leur exprime ici toute ma sincere reconnaissance.",
]
for d in deds:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(d)
    r.italic = True; r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = DARK; r.font.name = 'Times New Roman'

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# REMERCIEMENTS
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
h_rem = doc.add_heading('REMERCIEMENT', level=1)
for r in h_rem.runs:
    r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
h_rem.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

rems = [
    ("Je tiens a exprimer ma sincere gratitude envers toutes les personnes qui ont contribue "
     "au succes de ce projet de fin d'etudes portant sur le developpement du Portail de Gestion "
     "des Droits a la Retraite de la CIMR."),
    ("Je tiens a exprimer ma profonde gratitude envers mon encadrant pedagogique a l'EMSI, "
     "pour ses precieux conseils, sa disponibilite et son suivi rigoureux tout au long de "
     "la realisation de ce projet. Votre pedagogie et votre engagement pour notre reussite "
     "academique ont ete une source d'inspiration et d'encouragement constant."),
    ("Je remercie chaleureusement mon encadrant professionnel au sein de la CIMR pour "
     "son soutien, son encadrement et ses conseils avertis qui ont grandement enrichi "
     "mon experience et m'ont permis d'acquerir des competences essentielles dans le "
     "domaine du developpement logiciel et de la gestion des systemes d'information."),
    ("Je souhaite egalement adresser mes remerciements a l'ensemble de l'equipe technique "
     "de la CIMR qui m'a accueilli avec bienveillance et m'a permis de m'integrer "
     "rapidement dans l'environnement de travail. Votre collaboration et votre ouverture "
     "d'esprit ont ete inestimables pour la reussite de ce projet."),
    ("Je n'oublie pas de remercier les membres du jury pour avoir accepte d'evaluer et "
     "de juger ce travail. Leur participation constitue un honneur pour moi."),
    ("Enfin, je tiens a exprimer ma reconnaissance envers l'Ecole Marocaine des Sciences "
     "de l'Ingenieur (EMSI) pour m'avoir donne l'opportunite d'effectuer ce projet, "
     "qui a ete une experience formatrice et enrichissante. Ce projet de fin d'etudes a "
     "represente une etape importante de ma formation, et je suis reconnaissant envers "
     "chacun d'entre vous pour votre contribution a sa reussite."),
]
for txt in rems:
    body_para(doc, txt)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# RESUME
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hr = doc.add_heading('Resume', level=1)
for r in hr.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

body_para(doc, (
    "Le present rapport constitue le resultat d'un projet de fin d'etudes realise au cours "
    "de ma quatrieme annee academique a l'Ecole Marocaine des Sciences de l'Ingenieur (EMSI), "
    "en collaboration avec la Caisse Interprofessionnelle Marocaine de Retraite (CIMR). "
    "L'objectif central de ce projet etait de concevoir, developper et deployer un portail "
    "web complet de gestion des droits a la retraite, en utilisant un ensemble de technologies "
    "modernes et performantes."
))
body_para(doc, (
    "Ce projet a ete initie pour repondre a un besoin strategique de la CIMR : digitaliser "
    "et centraliser la gestion des affilies, des cotisations, des liquidations de retraite "
    "et des paiements de pension au sein d'une plateforme unique, securisee et evolutive. "
    "L'architecture adoptee est une architecture microservices basee sur Java Spring Boot "
    "pour le backend, React/TypeScript pour le frontend, et Apache Kafka pour la communication "
    "asynchrone entre les services."
))
body_para(doc, (
    "La premiere phase du projet a consiste en une analyse approfondie des besoins "
    "fonctionnels et non-fonctionnels de la CIMR. Cette etape a ete suivie d'une conception "
    "detaillee de l'architecture logicielle, incluant la modelisation UML des cas d'utilisation, "
    "des diagrammes de sequence et du diagramme de classes."
))
body_para(doc, (
    "L'equipe a ensuite procede au developpement de neuf microservices independants : "
    "auth-service, affiliation-service, contribution-service, liquidation-service, "
    "payment-service, reversion-service, admin-service, saga-orchestrator et api-gateway. "
    "Une attention particuliere a ete portee a la securite (JWT), a la robustesse "
    "(gestion des pannes Kafka) et a la qualite du code (corrections de bugs, "
    "coherence des modeles de donnees)."
))
body_para(doc, (
    "En conclusion, ce projet a abouti a la livraison d'une solution complete et "
    "fonctionnelle qui ameliore considerablement la gestion des droits a la retraite "
    "au sein de la CIMR. Ce rapport vise a decrire en detail les differentes etapes "
    "de ce projet, les choix technologiques effectues, les defis releves et les succes "
    "obtenus tout au long de cette experience enrichissante."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
ha = doc.add_heading('ABSTRACT', level=1)
for r in ha.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
ha.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

body_para(doc, (
    "This report reflects the outcome of a final year project undertaken during my fourth "
    "academic year at the Moroccan School of Engineering Sciences (EMSI), in collaboration "
    "with the Caisse Interprofessionnelle Marocaine de Retraite (CIMR). The main objective "
    "of this project was to design, develop and deploy a comprehensive web portal for "
    "managing pension rights, using a set of modern and high-performance technologies."
))
body_para(doc, (
    "This project was initiated to meet a strategic need of CIMR: to digitalize and centralize "
    "the management of affiliates, contributions, pension liquidations and pension payments "
    "within a single, secure and scalable platform. The adopted architecture is a microservices "
    "architecture based on Java Spring Boot for the backend, React/TypeScript for the "
    "frontend, and Apache Kafka for asynchronous inter-service communication."
))
body_para(doc, (
    "The first phase of the project consisted of a thorough analysis of CIMR's functional "
    "and non-functional requirements. This step was followed by a detailed design of the "
    "software architecture, including UML modeling of use cases, sequence diagrams and "
    "the class diagram."
))
body_para(doc, (
    "The team then proceeded to develop nine independent microservices, with particular "
    "attention paid to security (JWT authentication), robustness (Kafka fault handling) "
    "and code quality (bug fixes, data model consistency). The project resulted in the "
    "delivery of a complete and functional solution that significantly improves the "
    "management of pension rights at CIMR."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# TABLE DES MATIERES
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
htm = doc.add_heading('Table des matieres', level=1)
for r in htm.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
htm.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

toc_items = [
    ("INTRODUCTION GENERALE", "14"),
    ("Chapitre 1 : CONTEXTE GENERAL DU PROJET", "16"),
    ("  Introduction", "17"),
    ("  1. Organisme d'accueil — La CIMR", "17"),
    ("    1.1 Presentation de la CIMR", "17"),
    ("    1.2 Mission et valeurs", "18"),
    ("    1.3 Organisation interne", "19"),
    ("  2. PRESENTATION DU PROJET", "21"),
    ("    2.1 Perimetre du projet", "21"),
    ("    2.2 Presentation des utilisateurs cibles", "24"),
    ("  3. CADRE DU PROJET", "25"),
    ("  4. PROBLEMATIQUE", "26"),
    ("  5. SOLUTION PROPOSEE", "29"),
    ("    5.1 Architecture microservices", "29"),
    ("    5.2 Fonctionnalites cles", "30"),
    ("    5.3 Gestion de projet — Methodologie Agile/SCRUM", "31"),
    ("  Conclusion", "35"),
    ("Chapitre 2 : Analyse et Conception", "36"),
    ("  Introduction", "37"),
    ("  Le formalisme UML", "37"),
    ("  Analyse fonctionnelle — Diagrammes de cas d'utilisation", "38"),
    ("  Analyse dynamique — Diagrammes de sequence", "45"),
    ("  Analyse objet — Diagramme de classes", "51"),
    ("  Conclusion", "53"),
    ("Chapitre 3 : Etude Technique", "54"),
    ("  Introduction", "55"),
    ("  1. Environnement de developpement Backend", "55"),
    ("  2. Environnement de developpement Frontend", "62"),
    ("  3. Communication asynchrone — Apache Kafka", "67"),
    ("  4. Outils et Plateformes", "68"),
    ("  Conclusion", "72"),
    ("Chapitre 4 : Realisation", "73"),
    ("  Introduction", "74"),
    ("  Description des interfaces graphiques", "74"),
    ("  Conclusion", "82"),
    ("CONCLUSION ET PERSPECTIVES", "83"),
    ("WEBOGRAPHIE", "84"),
]

for title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    tabs = title.count('  ')
    p.paragraph_format.left_indent = Cm(tabs * 0.5)
    is_chapter = 'Chapitre' in title or title.isupper()
    r1 = p.add_run(title.strip())
    r1.bold = is_chapter
    r1.font.size = Pt(11 if is_chapter else 10.5)
    r1.font.color.rgb = DARK; r1.font.name = 'Times New Roman'
    dots_count = max(2, 60 - len(title.strip()) - tabs*2)
    r2 = p.add_run(' ' + '.' * dots_count + ' ')
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = 'Times New Roman'
    r3 = p.add_run(pg)
    r3.font.size = Pt(11); r3.font.color.rgb = DARK; r3.font.name = 'Times New Roman'

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# LISTE DES FIGURES
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hlf = doc.add_heading('LISTE DES FIGURES', level=1)
for r in hlf.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hlf.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

figures = [
    ("Figure 1", "Logo CIMR", "17"),
    ("Figure 2", "Secteur d'activite et partenaires CIMR", "18"),
    ("Figure 3", "Organigramme CIMR", "19"),
    ("Figure 4", "Chiffres cles CIMR", "20"),
    ("Figure 5", "Architecture globale du portail (microservices)", "22"),
    ("Figure 6", "Diagramme de Gantt du projet", "35"),
    ("Figure 7", "Diagramme de cas d'utilisation — Vue globale", "39"),
    ("Figure 8", "Diagramme de sequence — Authentification JWT", "46"),
    ("Figure 9", "Diagramme de sequence — Affiliation", "47"),
    ("Figure 10","Diagramme de sequence — Depot liquidation", "48"),
    ("Figure 11","Diagramme de sequence — Mise a jour statut", "49"),
    ("Figure 12","Diagramme de sequence — Notification Kafka", "50"),
    ("Figure 13","Diagramme de classes global", "52"),
    ("Figure 14","Logo Java", "55"),
    ("Figure 15","Logo Spring Boot", "56"),
    ("Figure 16","Architecture Spring Security + JWT", "57"),
    ("Figure 17","Logo React / TypeScript", "62"),
    ("Figure 18","Architecture frontend React (composants)", "63"),
    ("Figure 19","Logo Apache Kafka", "67"),
    ("Figure 20","Architecture Kafka Producer/Consumer", "67"),
    ("Figure 21","Logo IntelliJ IDEA", "68"),
    ("Figure 22","Logo Visual Studio Code", "68"),
    ("Figure 23","Logo Postman", "69"),
    ("Figure 24","Logo Docker / Kubernetes", "70"),
    ("Figure 25","Page de connexion", "74"),
    ("Figure 26","Tableau de bord administrateur", "75"),
    ("Figure 27","Gestion des affilies", "76"),
    ("Figure 28","Formulaire de depot liquidation", "77"),
    ("Figure 29","Suivi de dossier (vue affilie)", "78"),
    ("Figure 30","Page de gestion des liquidations (vue admin)", "79"),
    ("Figure 31","Centre de notifications", "80"),
    ("Figure 32","Simulation de pension", "81"),
    ("Figure 33","Structure output fichier SWIFT", "82"),
]

for fig_id, fig_title, pg in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{fig_id}: {fig_title} ")
    r1.font.size = Pt(11); r1.font.color.rgb = DARK; r1.font.name = 'Times New Roman'
    dots = '.' * max(2, 55 - len(f"{fig_id}: {fig_title}"))
    r2 = p.add_run(dots + ' ')
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = 'Times New Roman'
    r3 = p.add_run(pg)
    r3.font.size = Pt(11); r3.font.color.rgb = DARK; r3.font.name = 'Times New Roman'

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# LISTE DES TABLEAUX
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hlt = doc.add_heading('LISTE DES TABLEAUX', level=1)
for r in hlt.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hlt.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

tableaux = [
    ("Tableau 1", "Methode QQOQCP — Analyse du projet", "28"),
    ("Tableau 2", "Cas d'utilisation — Affiliation d'un nou vel affilie", "40"),
    ("Tableau 3", "Cas d'utilisation — Depot de demande de liquidation", "41"),
    ("Tableau 4", "Cas d'utilisation — Changement d'etat dossier", "42"),
    ("Tableau 5", "Cas d'utilisation — Consultation du livret individuel", "43"),
    ("Tableau 6", "Cas d'utilisation — Gestion des paiements", "44"),
    ("Tableau 7", "Cas d'utilisation — Gestion des utilisateurs", "45"),
    ("Tableau 8", "Stack technologique resume", "55"),
    ("Tableau 9", "Recapitulatif des bugs identifies et corriges", "72"),
]

for tab_id, tab_title, pg in tableaux:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{tab_id}: {tab_title} ")
    r1.font.size = Pt(11); r1.font.color.rgb = DARK; r1.font.name = 'Times New Roman'
    dots = '.' * max(2, 50 - len(f"{tab_id}: {tab_title}"))
    r2 = p.add_run(dots + ' ')
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = 'Times New Roman'
    r3 = p.add_run(pg)
    r3.font.size = Pt(11); r3.font.color.rgb = DARK; r3.font.name = 'Times New Roman'

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# LISTE DES ABREVIATIONS
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hla = doc.add_heading('LISTE DES ABREVIATIONS', level=1)
for r in hla.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hla.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

abrevs = [
    ("CIMR",  "Caisse Interprofessionnelle Marocaine de Retraite"),
    ("EMSI",  "Ecole Marocaine des Sciences de l'Ingenieur"),
    ("API",   "Application Programming Interface"),
    ("JWT",   "JSON Web Token"),
    ("REST",  "Representational State Transfer"),
    ("JSON",  "JavaScript Object Notation"),
    ("UML",   "Unified Modeling Language"),
    ("JPA",   "Java Persistence API"),
    ("ORM",   "Object-Relational Mapping"),
    ("MVC",   "Model-View-Controller"),
    ("SPA",   "Single Page Application"),
    ("CRUD",  "Create Read Update Delete"),
    ("IOC",   "Inversion Of Control"),
    ("TS",    "TypeScript"),
    ("SQL",   "Structured Query Language"),
    ("RGPD",  "Reglement General sur la Protection des Donnees"),
    ("CNDP",  "Commission Nationale de controle de la protection des Donnees a caractere Personnel"),
    ("CI/CD", "Continuous Integration / Continuous Deployment"),
    ("K8s",   "Kubernetes"),
    ("UUID",  "Universally Unique Identifier"),
]

tab_abr = doc.add_table(rows=len(abrevs), cols=2)
tab_abr.style = 'Table Grid'
for i, (abr, defi) in enumerate(abrevs):
    bg = 'EEF4FF' if i % 2 == 0 else 'FFFFFF'
    c0, c1 = tab_abr.rows[i].cells
    set_cell_bg(c0, bg); set_cell_bg(c1, bg)
    r0 = c0.paragraphs[0].add_run(abr)
    r0.bold = True; r0.font.size = Pt(11); r0.font.name = 'Times New Roman'
    r1 = c1.paragraphs[0].add_run(defi)
    r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
    for c in (c0, c1):
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after  = Pt(3)
    c0.width = Cm(3.5); c1.width = Cm(11.5)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# INTRODUCTION GENERALE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hi = doc.add_heading('INTRODUCTION GENERALE', level=1)
for r in hi.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hi.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_border_bottom(hi._p.getparent(), '000000', '6') if False else None
doc.add_paragraph()

body_para(doc, (
    "Dans un contexte economique et social en perpetuelle evolution, la gestion des "
    "systemes de retraite constitue un enjeu majeur pour les pays en voie de developpement. "
    "Au Maroc, la Caisse Interprofessionnelle Marocaine de Retraite (CIMR) joue un role "
    "fondamental dans la protection sociale des travailleurs du secteur prive, en garantissant "
    "la gestion et le versement de leurs droits a la retraite."
))
body_para(doc, (
    "Face a la croissance continue du nombre d'affilies et a la complexite grandissante "
    "des operations de gestion — affiliations, cotisations mensuelles, liquidations de "
    "dossiers de retraite, paiements de pension — la CIMR a ressenti la necessite imperative "
    "de digitaliser et de moderniser l'ensemble de ses processus internes. L'enjeu est "
    "double : ameliorer l'efficacite operationnelle des agents et offrir aux affilies un "
    "espace personnel transparent et accessible pour le suivi de leurs droits."
))
body_para(doc, (
    "C'est dans ce contexte qu'a ete initie le projet de developpement du Portail de Gestion "
    "des Droits a la Retraite CIMR. Ce projet ambitieux vise a concevoir et a deployer une "
    "plateforme web full-stack moderne, adoptant une architecture microservices scalable, "
    "capable de gerer l'integralite du cycle de vie d'un affilie : de son inscription jusqu'a "
    "la liquidation de ses droits et au versement de sa pension mensuelle."
))
body_para(doc, (
    "Le present rapport a pour objectif de presenter en detail le contexte et les motivations "
    "de ce projet, les etapes de sa planification et de son execution, ainsi que les resultats "
    "obtenus et les enseignements tires de cette experience professionnelle enrichissante. "
    "Au fil des chapitres qui suivent, nous explorerons les besoins de l'organisation, les "
    "choix d'architecture, les technologies adoptees, et les interfaces realisees."
))
body_para(doc, (
    "Ce rapport est egalement le fruit de mon stage de fin d'etudes au sein de la CIMR, "
    "d'une duree de plusieurs mois, realise dans le cadre de mon cursus academique en "
    "Ingenierie Informatique et Reseaux a l'EMSI. Ce stage m'a permis de mettre en "
    "pratique l'ensemble des competences acquises au cours de ma formation et d'acquerir "
    "une experience professionnelle solide dans le domaine du developpement logiciel."
))

# Structure du rapport
heading2(doc, "Structure du rapport")
body_para(doc, "Ce rapport est organise en quatre chapitres :")
items = [
    ("Chapitre 1", " presente le contexte general du projet : organisme d'accueil, presentation du projet, problematique et solution proposee."),
    ("Chapitre 2", " aborde l'analyse et la conception du systeme a travers les diagrammes UML."),
    ("Chapitre 3", " decrit l'environnement technique : langages, frameworks, outils et plateformes utilises."),
    ("Chapitre 4", " presente la realisation du projet a travers les interfaces graphiques developpees."),
]
for pref, txt in items:
    bullet_item(doc, txt, bold_prefix=pref)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 — PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
chapter_title_page(doc, "Chapitre 1", "CONTEXTE GENERAL DU PROJET")
page_break(doc)

# ── Chapitre 1 Contenu ──────────────────────────────────────────────────
add_logos_header(doc)

heading2(doc, "Introduction")
body_para(doc, (
    "L'objectif de ce premier chapitre est d'introduire le contexte general de notre projet "
    "de fin d'etudes. Nous commençons tout d'abord par la presentation de l'organisme "
    "d'accueil dans lequel s'est deroule notre stage, la CIMR. Nous exposons ensuite "
    "le cadre du projet en decrivant les besoins identifies, la problematique rencontree "
    "et la solution proposee pour y repondre. Nous presentons enfin la demarche de "
    "gestion de projet adoptee pour mener a bien cette initiative."
))

heading1(doc, "1.  Organisme d'accueil")
heading2(doc, "1.1  Presentation de la CIMR")
body_para(doc, (
    "La Caisse Interprofessionnelle Marocaine de Retraite (CIMR) est un organisme de "
    "retraite de droit prive marocain cree en 1949. Elle gere le regime de retraite "
    "complementaire des salaries du secteur prive au Maroc. Avec plus de 70 ans "
    "d'experience dans le domaine de la protection sociale, la CIMR est aujourd'hui "
    "l'un des piliers du systeme de retraite marocain."
))
body_para(doc, (
    "La CIMR compte plus de 1,2 million d'affilies actifs et gere plusieurs centaines "
    "de milliers de dossiers de retraite en cours de versement. Son action s'inscrit "
    "dans une demarche de service public, visant a garantir la securite financiere "
    "des travailleurs a l'age de la retraite."
))
figure_caption(doc, "Figure 1: Logo CIMR")

heading2(doc, "1.2  Mission et valeurs")
body_para(doc, "La CIMR s'articule autour de trois missions fondamentales :")
bullet_item(doc, "Collecter les cotisations salariales et patronales des entreprises adherentes et de leurs employes affilies ;")
bullet_item(doc, "Gerer les droits acquis (points de retraite) de chaque affilie tout au long de sa carriere professionnelle ;")
bullet_item(doc, "Liquider et verser les pensions de retraite aux affilies ayant atteint l'age de depart a la retraite ou remplissant les conditions requises.")
body_para(doc, (
    "La CIMR fonde son action sur des valeurs fortes : equite dans le traitement des "
    "beneficiaires, transparence dans la gestion des fonds, fiabilite dans le versement "
    "des droits et innovation dans ses pratiques et outils de gestion."
))
figure_caption(doc, "Figure 2: Secteur d'activite et partenaires CIMR")

heading2(doc, "1.3  Organisation interne")
body_para(doc, (
    "La CIMR est administree par un Conseil d'Administration paritaire compose de "
    "representants des employeurs et des salaries. Elle est structuree en plusieurs "
    "directions fonctionnelles : Direction des Affilies et des Adherents, Direction "
    "des Placements, Direction des Systemes d'Information, Direction des Ressources "
    "Humaines et Direction Financiere."
))
body_para(doc, (
    "Dans le cadre de ce projet, nous avons ete integres au sein de la Direction des "
    "Systemes d'Information (DSI), plus particulierement dans l'equipe chargee du "
    "developpement et de la maintenance des applications metier de la CIMR."
))
figure_caption(doc, "Figure 3: Organigramme CIMR")
figure_caption(doc, "Figure 4: Chiffres cles CIMR")

heading1(doc, "2.  PRESENTATION DU PROJET")
heading2(doc, "2.1  Perimetre du projet")

body_para(doc, (
    "Le projet de developpement du Portail de Gestion des Droits a la Retraite CIMR "
    "a ete initie pour repondre a un besoin crucial de modernisation des outils de "
    "gestion interne. Le perimetre de ce projet est defini comme suit :"
))

heading3(doc, "ANALYSE ET PLANIFICATION")
bullet_item(doc, "Analyse approfondie des processus de gestion des affilies, cotisations et liquidations existants.")
bullet_item(doc, "Identification des besoins specifiques et des exigences fonctionnelles et techniques.")
bullet_item(doc, "Planification des ressources humaines et techniques necessaires au projet.")
bullet_item(doc, "Etablissement d'un calendrier de projet incluant les etapes cles et les echeances.")

heading3(doc, "CONCEPTION ET DEVELOPPEMENT")
bullet_item(doc, "Conception d'une architecture microservices distribuee avec Spring Boot, React/TypeScript et Apache Kafka.")
bullet_item(doc, "Developpement de neuf microservices independants couvrant l'ensemble des fonctionnalites metier.")
bullet_item(doc, "Implementation d'un mecanisme d'authentification securise par tokens JWT.")
bullet_item(doc, "Integration d'un agent IA (Python/FastAPI) pour l'assistance aux affilies et la verification d'identite.")

heading3(doc, "INTEGRATION ET TEST")
bullet_item(doc, "Integration des microservices via un API Gateway centralise.")
bullet_item(doc, "Realisation de tests fonctionnels et de tests de regression sur les API REST.")
bullet_item(doc, "Validation de la coherence des modeles de donnees entre le frontend et le backend.")
bullet_item(doc, "Identification et correction des anomalies (9 bugs corriges lors de l'audit final).")

heading3(doc, "FORMATION ET DEPLOIEMENT")
bullet_item(doc, "Formation des utilisateurs finaux (agents CIMR et administrateurs) sur l'utilisation de la plateforme.")
bullet_item(doc, "Preparation du deploiement en environnement de production via Docker et Kubernetes.")
bullet_item(doc, "Transition progressive depuis les anciens outils vers la nouvelle plateforme.")

heading3(doc, "SUIVI ET MAINTENANCE")
bullet_item(doc, "Mise en place d'un monitoring des services via Spring Boot Actuator.")
bullet_item(doc, "Gestion des incidents et resolution des anomalies post-deploiement.")
bullet_item(doc, "Evolution continue de la solution en fonction des retours des utilisateurs.")

figure_caption(doc, "Figure 5: Architecture globale du portail (microservices)")

heading2(doc, "2.2  Presentation des utilisateurs cibles")
body_para(doc, "Le portail s'adresse a trois categories d'utilisateurs :")
bullet_item(doc, "Administrateurs CIMR : gestion complete des affilies, des cotisations, des liquidations et des paiements ;", bold_prefix="")
bullet_item(doc, "Agents CIMR : traitement des dossiers et suivi des demandes des affilies ;")
bullet_item(doc, "Affilies : acces a leur espace personnel pour le suivi de leurs droits, le depot de demandes de liquidation et la consultation de leurs bulletins.")

heading1(doc, "3.  CADRE DU PROJET")
body_para(doc, (
    "Ce projet s'inscrit dans le cadre de la transformation numerique de la CIMR. "
    "Il vise a mettre en place une solution informatique innovante permettant de :"
))
bullet_item(doc, "Centraliser l'ensemble des donnees des affilies dans une base de donnees structuree et securisee ;")
bullet_item(doc, "Automatiser les calculs de points de retraite selon les regles actuarielles de la CIMR (Article 6) ;")
bullet_item(doc, "Simplifier et accelerer le processus de depot et de traitement des dossiers de liquidation de retraite ;")
bullet_item(doc, "Garantir la traçabilite de toutes les operations grace a un systeme d'audit logs exhaustif ;")
bullet_item(doc, "Assurer la conformite RGPD et CNDP dans le traitement des donnees personnelles des affilies.")

heading1(doc, "4.  PROBLEMATIQUE")
body_para(doc, (
    "La gestion des droits a la retraite au sein de la CIMR est actuellement confrontee "
    "a plusieurs defis et inefficacites qui necessitent une solution adaptee. Plusieurs "
    "problematiques cles se posent :"
))
bullet_item(doc, "Dispersion des donnees : les informations des affilies sont dispersees entre plusieurs systemes heterogenes, rendant difficile l'obtention d'une vue unifiee du dossier d'un affilie ;", bold_prefix="Dispersion des donnees : ")
bullet_item(doc, "Processus manuels : de nombreuses etapes du traitement des dossiers (calcul des cotisations, validation des liquidations) sont encore realisees manuellement, ce qui augmente les risques d'erreurs et rallonge les delais de traitement ;", bold_prefix="")
bullet_item(doc, "Manque de transparence : les affilies n'ont pas acces en temps reel a l'etat de leurs droits acquis ni a l'avancement de leurs dossiers en cours de traitement ;", bold_prefix="")
bullet_item(doc, "Securite insuffisante : les systemes existants ne disposent pas de mecanismes d'authentification modernes ni de controle d'acces granulaire par role d'utilisateur ;", bold_prefix="")
bullet_item(doc, "Absence de notifications : aucun systeme de notification automatique n'informe les affilies ou les agents des changements d'etat de leurs dossiers.", bold_prefix="")

table_caption(doc, "Tableau 1: Methode QQOQCP — Analyse du projet")
simple_table(doc,
    headers=["", "Question", "Reponse"],
    rows=[
        ["Qui ?",     "Qui est concerne par ce projet ?",
         "La CIMR, ses agents et ses affilies"],
        ["Quoi ?",    "Quelles sont les principales inefficacites ?",
         "Processus manuels, donnees dispersees, absence de traçabilite"],
        ["Ou ?",      "Ou se deroule la gestion des dossiers ?",
         "Au sein des directions metier de la CIMR"],
        ["Quand ?",   "Quels sont les delais du projet ?",
         "Definis dans le diagramme de Gantt (6 mois)"],
        ["Comment ?", "Comment les donnees sont-elles traitees actuellement ?",
         "Via des fichiers Excel et des logiciels legacy non interconnectes"],
        ["Pourquoi ?","Pourquoi moderniser le systeme ?",
         "Reduire les erreurs, accelerer le traitement et ameliorer l'experience affilie"],
    ],
    col_widths=[2.0, 5.5, 7.5]
)

heading1(doc, "5.  SOLUTION PROPOSEE")
heading2(doc, "5.1  Architecture microservices")
bullet_item(doc, "API Gateway (Spring Cloud Gateway) : point d'entree unique, routage des requetes, gestion CORS ;", bold_prefix="API Gateway : ")
bullet_item(doc, "Auth Service : authentification JWT, gestion des utilisateurs, reinitialisation de mot de passe ;", bold_prefix="Auth Service : ")
bullet_item(doc, "Affiliation Service : CRUD affilies, bulletins, justificatifs, conformite RGPD/CNDP ;", bold_prefix="Affiliation Service : ")
bullet_item(doc, "Contribution Service : cotisations Article 6, ledger de points, simulation actuarielle, achat de points ;", bold_prefix="Contribution Service : ")
bullet_item(doc, "Liquidation Service : depot de dossiers, upload de documents, suivi d'etat, notifications Kafka ;", bold_prefix="Liquidation Service : ")
bullet_item(doc, "Payment Service : allocations et paiements de pension ;", bold_prefix="Payment Service : ")
bullet_item(doc, "Reversion Service : gestion des droits des ayants droit ;", bold_prefix="Reversion Service : ")
bullet_item(doc, "Admin Service : audit logs, CNDP, notifications, tickets support ;", bold_prefix="Admin Service : ")
bullet_item(doc, "AI Agent Service (Python/FastAPI) : chatbot support, verification CIN par OCR.", bold_prefix="AI Agent Service : ")

heading2(doc, "5.2  Fonctionnalites cles")
bullet_item(doc, "Tableau de bord administrateur avec statistiques en temps reel et alertes ;")
bullet_item(doc, "Gestion complete du cycle de vie des affilies : affiliation, radiation, suspension ;")
bullet_item(doc, "Import CSV massif des cotisations avec validation et detection des doublons ;")
bullet_item(doc, "Simulation actuarielle de pension basee sur les parametres de l'Article 6 CIMR ;")
bullet_item(doc, "Depot en ligne des demandes de liquidation avec upload de documents ;")
bullet_item(doc, "Systeme de notifications en quasi-temps-reel via Apache Kafka ;")
bullet_item(doc, "Support multilingue (francais/arabe) avec gestion RTL.")

heading2(doc, "5.3  Gestion de projet — Methodologie Agile/SCRUM")
body_para(doc, (
    "Pour la gestion de ce projet, nous avons adopte la methodologie Agile avec le "
    "cadre SCRUM. Ce choix se justifie par la nature iterative du projet, la necessite "
    "d'adapter continuellement les fonctionnalites aux retours des utilisateurs et "
    "la volonte de livrer de la valeur de maniere incrementale."
))

heading3(doc, "SCRUM comme fil conducteur")
body_para(doc, (
    "La methodologie SCRUM repose sur des iterations courtes (sprints de 2 semaines) "
    "permettant de livrer regularierement des fonctionnalites testables. Chaque sprint "
    "demarre par une seance de planification et se termine par une revue et une "
    "retrospective. Le backlog produit a ete maintenu et mis a jour tout au long du projet."
))

heading3(doc, "Roles et rituels")
bullet_item(doc, "Product Owner : responsable de la vision produit et de la priorisation du backlog ;")
bullet_item(doc, "Scrum Master : garant de l'application de la methode et de la productivite de l'equipe ;")
bullet_item(doc, "Equipe de developpement : responsable de la realisation des fonctionnalites.")
body_para(doc, (
    "Les rituels SCRUM pratiques incluaient le stand-up quotidien (15 min), la "
    "planification de sprint, la revue de sprint et la retrospective."
))

heading3(doc, "Diagramme de Gantt")
body_para(doc, (
    "Le diagramme de Gantt represente la planification temporelle du projet, "
    "decoupee en phases : collecte des besoins, modelisation des processus, "
    "selection des technologies, conception applicative, conception technique, "
    "realisation et tests, integration et livraison."
))
figure_caption(doc, "Figure 6: Diagramme de Gantt du projet")

heading2(doc, "Conclusion")
body_para(doc, (
    "Apres avoir expose le contexte general du projet, en particulier la presentation "
    "de l'organisme d'accueil CIMR, la problematique identifiee et la solution proposee "
    "avec la methodologie Agile/SCRUM, il convient desormais de passer a l'etape "
    "d'analyse et de conception de la solution."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 — PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
chapter_title_page(doc, "Chapitre 2", "Analyse et Conception")
page_break(doc)

# ── Chapitre 2 Contenu ──────────────────────────────────────────────────
add_logos_header(doc)

heading2(doc, "Introduction")
body_para(doc, (
    "Dans ce chapitre, nous allons aborder l'analyse et la conception du systeme. "
    "Nous commencerons par une analyse des besoins et des exigences pour comprendre "
    "les fonctionnalites attendues du systeme. Ensuite, nous presenterons la conception "
    "du systeme en utilisant des diagrammes UML pour modeliser les differents aspects "
    "du systeme. Cette etape est cruciale pour garantir que le systeme repond aux "
    "besoins et aux exigences identifies lors de l'analyse."
))

heading1(doc, "Le formalisme UML")
body_para(doc, (
    "UML (Unified Modeling Language) est une notation permettant la modelisation "
    "d'un probleme. Ce langage est ne de la fusion de plusieurs methodes existantes "
    "et est devenu la reference en termes de modelisation objet. Dans le cadre de "
    "ce projet, nous avons utilise quatre types de diagrammes UML :"
))
bullet_item(doc, "Les diagrammes de cas d'utilisation : representation des fonctions du systeme du point de vue de l'utilisateur ;")
bullet_item(doc, "Les diagrammes de sequence : representation temporelle des objets et de leurs interactions ;")
bullet_item(doc, "Les diagrammes d'activites : representation du comportement d'un processus metier ;")
bullet_item(doc, "Le diagramme de classes : representation de la structure statique en termes de classes et de relations.")

heading1(doc, "Analyse fonctionnelle")
heading2(doc, "Diagramme de cas d'utilisation")
body_para(doc, (
    "L'objectif principal du systeme est de gerer l'ensemble du cycle de vie d'un "
    "affilie CIMR, depuis son inscription jusqu'a la liquidation de ses droits. "
    "Les acteurs identifies sont :"
))
bullet_item(doc, "Administrateur : acces complet a toutes les fonctionnalites du systeme ;")
bullet_item(doc, "Agent CIMR : gestion des affilies et traitement des dossiers de liquidation ;")
bullet_item(doc, "Affilie : consultation de ses droits, depot de demandes et suivi de dossiers.")
figure_caption(doc, "Figure 7: Diagramme de cas d'utilisation — Vue globale")

# Tableau CU1
table_caption(doc, "Tableau 2: Cas d'utilisation — Affiliation d'un nouvel affilie")
use_case_table(doc,
    titre    = "Cas d'utilisation : Affilier un nouveau salarie",
    but      = "Enregistrer un nouvel affilie dans le systeme CIMR",
    acteur   = "Administrateur / Agent",
    pre      = "L'acteur doit etre authentifie avec le role ROLE_ADMIN ou ROLE_AGENT.",
    scenario_p = [
        "L'agent accede au module 'Affilies' et clique sur 'Nouvel affilie'.",
        "Le systeme affiche le formulaire de creation.",
        "L'agent saisit les informations personnelles, professionnelles et les pieces justificatives.",
        "L'agent soumet le formulaire.",
        "Le systeme valide les donnees et cree le compte affilie et l'acces utilisateur.",
        "Le systeme genere un numero d'immatriculation unique et envoie les identifiants par email.",
    ],
    scenario_a = [
        "Si le CIN est deja enregistre, le systeme affiche un message d'erreur de doublon.",
        "Si les champs obligatoires sont manquants, le systeme affiche les erreurs de validation.",
    ]
)

# Tableau CU2
table_caption(doc, "Tableau 3: Cas d'utilisation — Depot de demande de liquidation")
use_case_table(doc,
    titre    = "Cas d'utilisation : Deposer une demande de liquidation",
    but      = "Permettre a un affilie de deposer sa demande de mise en retraite",
    acteur   = "Affilie",
    pre      = "L'affilie doit etre authentifie. Il ne doit pas avoir de demande en cours.",
    scenario_p = [
        "L'affilie accede au module 'Liquidation' et clique sur 'Deposer ma demande'.",
        "Le systeme affiche le formulaire avec la liste des documents requis.",
        "L'affilie televerse les copies numeriques (CIN, attestation, RIB, acte de naissance).",
        "L'affilie confirme avoir lu les conditions et l'obligation de se presenter en agence.",
        "L'affilie confirme et soumet sa demande.",
        "Le systeme enregistre la demande et notifie les administrateurs via Kafka.",
    ],
    scenario_a = [
        "Si un document depasse 5 Mo, le systeme refuse le telechargement.",
        "Si l'identifiant de l'affilie est manquant, le systeme bloque la soumission.",
    ]
)

# Tableau CU3
table_caption(doc, "Tableau 4: Cas d'utilisation — Changement d'etat d'un dossier")
use_case_table(doc,
    titre    = "Cas d'utilisation : Changer l'etat d'un dossier de liquidation",
    but      = "Permettre a un administrateur de faire avancer le traitement d'un dossier",
    acteur   = "Administrateur / Agent",
    pre      = "L'administrateur doit etre authentifie. Le dossier doit exister.",
    scenario_p = [
        "L'administrateur accede a la liste des dossiers en attente.",
        "Il selectionne un dossier et clique sur 'Changer le statut'.",
        "Le systeme affiche le formulaire de changement d'etat.",
        "L'administrateur choisit le nouvel etat et saisit un commentaire si necessaire.",
        "Le systeme valide, enregistre et notifie l'affilie par Kafka.",
    ],
    scenario_a = [
        "Si aucun commentaire n'est saisi pour un rejet, le systeme affiche une erreur.",
    ]
)

# Tableau CU4
table_caption(doc, "Tableau 5: Cas d'utilisation — Consultation du livret individuel")
use_case_table(doc,
    titre    = "Cas d'utilisation : Consulter le livret individuel de points",
    but      = "Permettre a un affilie de consulter son historique de cotisations et de points",
    acteur   = "Affilie / Administrateur",
    pre      = "L'utilisateur doit etre authentifie.",
    scenario_p = [
        "L'affilie accede au module 'Cotisations'.",
        "Le systeme recupere l'historique des cotisations et le ledger de points.",
        "Le systeme affiche la synthese : total des points, estimation de pension mensuelle.",
        "L'affilie peut consulter le detail par periode et par contribution.",
    ],
    scenario_a = [
        "Si aucune cotisation n'est enregistree, le systeme affiche un message informatif.",
    ]
)

heading1(doc, "Analyse dynamique")
body_para(doc, (
    "Les diagrammes de sequence permettent de representer l'ordre chronologique "
    "des interactions entre les differents acteurs et composants du systeme pour "
    "chaque cas d'utilisation identifie."
))
figure_caption(doc, "Figure 8: Diagramme de sequence — Authentification JWT")
body_para(doc, (
    "Le diagramme de sequence d'authentification montre comment un utilisateur se "
    "connecte au systeme. Il detaille les etapes de saisie des identifiants, de "
    "verification dans la base de donnees, et de generation d'un token JWT en cas "
    "de succes. Le token est ensuite utilise pour authentifier toutes les requetes "
    "subsequentes via le header Authorization."
))
figure_caption(doc, "Figure 9: Diagramme de sequence — Affiliation")
figure_caption(doc, "Figure 10: Diagramme de sequence — Depot liquidation")
figure_caption(doc, "Figure 11: Diagramme de sequence — Mise a jour statut")
figure_caption(doc, "Figure 12: Diagramme de sequence — Notification Kafka")
body_para(doc, (
    "Le diagramme de sequence de notification Kafka illustre le mecanisme de "
    "communication asynchrone entre le Liquidation Service (producteur) et l'Admin "
    "Service (consommateur). Lors de la creation d'une demande, le Liquidation Service "
    "publie un evenement dans le topic 'notifications-topic'. L'Admin Service consomme "
    "cet evenement et persiste la notification en base de donnees, la rendant accessible "
    "via l'API de notifications."
))

heading1(doc, "Analyse objet")
body_para(doc, (
    "Le diagramme de classes represente l'aspect statique du systeme en termes de "
    "classes et des relations entre elles. Il identifie les attributs et les methodes "
    "qui refletent respectivement les proprietes et les services offerts par notre systeme."
))
figure_caption(doc, "Figure 13: Diagramme de classes global")
body_para(doc, (
    "Les entites principales identifiees sont : Affilie, BulletinAffiliation, Justificatif, "
    "Radiation, Contribution, PointsLedger, PointsPurchase, DemandeLiquidation, "
    "DossierDocument, Paiement, Allocation, AyantDroit, Notification, AuditLog, User, "
    "PasswordResetToken et SupportTicket."
))

heading2(doc, "Conclusion")
body_para(doc, (
    "Dans ce chapitre, nous avons presente l'etude preliminaire et l'analyse "
    "fonctionnelle du projet. Nous avons identifie les acteurs et les interactions "
    "entre les acteurs et le systeme, modelise les cas d'utilisation, realise les "
    "diagrammes de sequence et extrait les entites du systeme via le diagramme de "
    "classes. Ces elements constituent la base solide sur laquelle repose la "
    "conception technique du portail."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 — PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
chapter_title_page(doc, "Chapitre 3", "Etude Technique")
page_break(doc)

# ── Chapitre 3 Contenu ──────────────────────────────────────────────────
add_logos_header(doc)

heading2(doc, "Introduction")
body_para(doc, (
    "Ce chapitre est consacre a la partie technique du projet. Apres avoir termine "
    "la phase de specification et de conception, nous presentons ici les choix "
    "technologiques effectues, l'environnement de developpement adopte et les outils "
    "utilises pour la realisation de ce portail. Ces choix ont ete guides par les "
    "exigences de performance, de scalabilite, de securite et de maintenabilite du systeme."
))

heading1(doc, "1.  Environnement de developpement Backend")
heading2(doc, "1.1  Java")
body_para(doc, (
    "Java est un langage de programmation oriente objet cree en 1995. Sa philosophie "
    "'Write Once, Run Anywhere' repose sur la compilation vers un bytecode execute "
    "dans une machine virtuelle Java (JVM). Pour ce projet, nous avons utilise "
    "Java 21 (LTS), qui apporte des ameliorations significatives en termes de "
    "performances et de nouvelles fonctionnalites du langage."
))
figure_caption(doc, "Figure 14: Logo Java")

heading2(doc, "1.2  Spring Boot")
body_para(doc, (
    "Spring Boot est un framework Java qui facilite la creation d'applications "
    "web et de microservices. Il inclut un serveur d'application embarque (Tomcat), "
    "une autoconfiguration intelligente et un ecosysteme de modules (Spring Data, "
    "Spring Security, Spring Cloud). Sa version 3.x, basee sur Jakarta EE 10, "
    "a ete adoptee pour sa compatibilite avec Java 21 et ses performances ameliorees."
))
bullet_item(doc, "Serveur embarque Tomcat — pas de deploiement WAR necessaire ;")
bullet_item(doc, "Spring Data JPA — abstraction ORM pour la persistence des donnees ;")
bullet_item(doc, "Spring Cloud Gateway — routage et filtrage des requetes API ;")
bullet_item(doc, "Actuator — monitoring et health checks des services.")
figure_caption(doc, "Figure 15: Logo Spring Boot")

heading2(doc, "1.3  Spring Security et JWT")
body_para(doc, (
    "Spring Security est le framework de securite standard pour les applications "
    "Spring. Dans ce projet, il est configure en mode stateless avec une "
    "authentification par tokens JWT (JSON Web Token). Chaque token contient les "
    "informations d'identite (username, roles) et est signe avec une cle secrete "
    "pour garantir son integrite."
))
bullet_item(doc, "Authentification : verification du username/mot de passe, generation du JWT ;")
bullet_item(doc, "Autorisation : verification du token a chaque requete, controle des roles ;")
bullet_item(doc, "BCrypt : hachage des mots de passe avec un facteur de cout eleve.")
figure_caption(doc, "Figure 16: Architecture Spring Security + JWT")

heading2(doc, "1.4  Apache Kafka")
body_para(doc, (
    "Apache Kafka est une plateforme de streaming distribue utilisee pour la "
    "communication asynchrone entre les microservices. Dans ce projet, Kafka est "
    "utilise pour la transmission des evenements de notification entre le Liquidation "
    "Service (producteur) et l'Admin Service (consommateur)."
))
bullet_item(doc, "Topic 'notifications-topic' : evenements de liquidation et de mise a jour de statut ;")
bullet_item(doc, "Producer : LiquidationService — envoie les evenements de notification ;")
bullet_item(doc, "Consumer : NotificationConsumer (Admin Service) — persiste les notifications.")
figure_caption(doc, "Figure 19: Logo Apache Kafka")
figure_caption(doc, "Figure 20: Architecture Kafka Producer/Consumer")

heading1(doc, "2.  Environnement de developpement Frontend")
heading2(doc, "2.1  TypeScript")
body_para(doc, (
    "TypeScript est un sur-ensemble type de JavaScript developpe par Microsoft. "
    "Il introduit le typage statique optionnel, les interfaces, les generiques et "
    "la programmation orientee objet. Le code TypeScript est transpile en JavaScript "
    "avant execution dans le navigateur. L'utilisation de TypeScript permet de detecter "
    "les erreurs de type a la compilation plutot qu'a l'execution."
))

heading2(doc, "2.2  React 19 et l'ecosysteme frontend")
body_para(doc, (
    "React est une bibliotheque JavaScript developpee par Meta pour la creation "
    "d'interfaces utilisateur. Sa version 19 apporte des ameliorations de performances "
    "et de nouvelles fonctionnalites hooks. L'ecosysteme frontend adopte dans ce projet comprend :"
))
bullet_item(doc, "React 19 — bibliotheque UI avec composants fonctionnels et hooks ;")
bullet_item(doc, "Vite 8 — outil de build ultra-rapide avec Hot Module Replacement ;")
bullet_item(doc, "React Router DOM 7 — routage client SPA ;")
bullet_item(doc, "React Hook Form 7 + Zod 4 — gestion et validation des formulaires ;")
bullet_item(doc, "Axios 1.x — client HTTP avec intercepteurs JWT automatiques ;")
bullet_item(doc, "Framer Motion 12 — animations fluides ;")
bullet_item(doc, "React Hot Toast — notifications utilisateur.")
figure_caption(doc, "Figure 17: Logo React / TypeScript")
figure_caption(doc, "Figure 18: Architecture frontend React (composants)")

heading1(doc, "3.  Outils et Plateformes")
heading2(doc, "3.1  IntelliJ IDEA")
body_para(doc, (
    "IntelliJ IDEA est l'IDE developpe par JetBrains, specialise pour le "
    "developpement Java. Il offre une assistance intelligente au code, un debogueur "
    "puissant et une integration native avec Spring Boot."
))
figure_caption(doc, "Figure 21: Logo IntelliJ IDEA")

heading2(doc, "3.2  Visual Studio Code")
body_para(doc, (
    "VS Code est l'editeur de code open-source developpe par Microsoft. Leger et "
    "performant, il a ete utilise pour le developpement du frontend React/TypeScript "
    "et de l'agent IA Python, grace a ses nombreuses extensions (ESLint, Prettier, "
    "Python, REST Client)."
))
figure_caption(doc, "Figure 22: Logo Visual Studio Code")

heading2(doc, "3.3  Postman")
body_para(doc, (
    "Postman est un outil de test d'API REST. Il a ete utilise pour tester "
    "et documenter tous les endpoints des microservices, en organisant les "
    "requetes en collections par service."
))
figure_caption(doc, "Figure 23: Logo Postman")

heading2(doc, "3.4  Docker et Kubernetes")
body_para(doc, (
    "Docker permet la containerisation des microservices, garantissant la portabilite "
    "et la reproductibilite de l'environnement d'execution. Kubernetes (K8s) assure "
    "l'orchestration des conteneurs en production, avec la gestion de la scalabilite "
    "horizontale et du self-healing des services."
))
figure_caption(doc, "Figure 24: Logo Docker / Kubernetes")

heading2(doc, "3.5  Recapitulatif technologique")
table_caption(doc, "Tableau 8: Stack technologique resume")
simple_table(doc,
    headers=["Couche", "Technologie", "Version", "Role"],
    rows=[
        ["Backend",   "Java",              "21 (LTS)", "Langage principal"],
        ["Backend",   "Spring Boot",       "3.x",      "Framework microservices"],
        ["Backend",   "Spring Security",   "6.x",      "Securite / JWT"],
        ["Backend",   "Spring Data JPA",   "3.x",      "Persistance ORM"],
        ["Backend",   "Apache Kafka",      "3.x",      "Messaging asynchrone"],
        ["Backend",   "Lombok",            "latest",   "Reduction boilerplate"],
        ["Frontend",  "React",             "19.x",     "Interface utilisateur"],
        ["Frontend",  "TypeScript",        "5.9.x",    "Typage statique"],
        ["Frontend",  "Vite",              "8.x",      "Build / dev server"],
        ["Frontend",  "React Router",      "7.x",      "Routage SPA"],
        ["Frontend",  "Zod",               "4.x",      "Validation schemas"],
        ["Frontend",  "Axios",             "1.x",      "Client HTTP"],
        ["IA",        "Python / FastAPI",  "3.12",     "Agent IA / OCR CIN"],
        ["Infra",     "Docker",            "latest",   "Containerisation"],
        ["Infra",     "Kubernetes",        "latest",   "Orchestration"],
    ],
    col_widths=[2.5, 4.0, 3.0, 5.5]
)

heading2(doc, "Bilan audit qualite — Bugs identifies et corriges")
table_caption(doc, "Tableau 9: Recapitulatif des bugs identifies et corriges")
simple_table(doc,
    headers=["#", "Fichier", "Severite", "Description", "Correction"],
    rows=[
        ["1", "Notification.java",       "CRITIQUE", "isRead serialise en 'read' (Jackson)",  "@JsonProperty('isRead')"],
        ["2", "DemandeLiquidation.java", "HIGH",     "3 champs manquants en BDD",             "Ajout enum + champs JPA"],
        ["3", "DossierDocument.java",    "HIGH",     "4 champs metadonnees manquants",        "Ajout nomFichier, taille, date, statut"],
        ["4", "LiquidationService.java", "HIGH",     "Kafka sans try/catch",                  "Try/catch + log.warn"],
        ["5", "LiquidationService.java", "HIGH",     "Metadonnees doc non persistees",        "Renseignement depuis MultipartFile"],
        ["6", "LiquidationService.java", "MEDIUM",   "probeContentType retourne null",        "Null check avant MediaType"],
        ["7", "liquidations.ts",         "MEDIUM",   "Valeurs hardcodees dans mapping",       "Lecture champs reels backend"],
        ["8", "liquidations.ts",         "MEDIUM",   "typeLiquidation non envoye",            "Ajout dans payload create()"],
        ["9", "LiquidationFormPage.tsx", "MEDIUM",   "Pas validation affilieId vide",         "Guard + toast erreur"],
    ],
    col_widths=[0.8, 4.5, 2.0, 4.8, 4.0]
)

heading2(doc, "Conclusion")
body_para(doc, (
    "Dans ce chapitre, nous avons presente les choix technologiques du projet et "
    "l'ensemble de l'environnement de developpement utilise. La combinaison de "
    "Spring Boot pour le backend, React/TypeScript pour le frontend et Kafka pour "
    "la communication asynchrone constitue une architecture moderne et robuste. "
    "Nous avons egalement presente le bilan de l'audit qualite effectue sur la "
    "codebase, ayant abouti a la correction de 9 bugs dont 2 de severite critique."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# CHAPITRE 4 — PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
chapter_title_page(doc, "Chapitre 4", "Realisation")
page_break(doc)

# ── Chapitre 4 Contenu ──────────────────────────────────────────────────
add_logos_header(doc)

heading2(doc, "Introduction")
body_para(doc, (
    "Le present chapitre constitue la derniere phase de la mise en place de la solution. "
    "Il est consacre a la description des interfaces graphiques realisees ainsi qu'a "
    "la presentation des fonctionnalites implementees. Chaque interface est presentee "
    "avec une description de son fonctionnement et de sa valeur ajoutee pour l'utilisateur."
))

heading1(doc, "Description des interfaces graphiques")

heading2(doc, "Page de connexion")
body_para(doc, (
    "La page de connexion est le point d'entree de l'application. Elle propose une "
    "interface intuitive avec un formulaire de saisie des identifiants, une option "
    "de recuperation de mot de passe (par email ou par verification CIN via l'agent IA), "
    "et un formulaire de contact pour le support. Une animation de succes confirme "
    "visuellement la connexion avant la redirection vers le tableau de bord."
))
figure_caption(doc, "Figure 25: Page de connexion")

heading2(doc, "Tableau de bord administrateur")
body_para(doc, (
    "Le tableau de bord administrateur offre une vue d'ensemble de l'activite CIMR "
    "en temps reel. Il affiche quatre indicateurs cles : nombre d'affilies actifs, "
    "montant des cotisations du mois en cours, nombre de dossiers en attente et "
    "paiements planifies. Une table des dossiers recents avec leur priorite "
    "(Haute/Moyenne/Basse) permet un traitement rapide. Les tickets de support "
    "reçus sont egalement affiches avec la possibilite d'envoyer des emails de "
    "reinitialisation de mot de passe directement depuis le tableau de bord."
))
figure_caption(doc, "Figure 26: Tableau de bord administrateur")

heading2(doc, "Gestion des affilies")
body_para(doc, (
    "Le module de gestion des affilies permet aux administrateurs et agents de "
    "creer, modifier, suspendre ou radier des affilies. Le formulaire de creation "
    "inclut la validation par schema Zod et la gestion du consentement CNDP. "
    "Un systeme d'import CSV massif permet d'integrer des cotisations en lot "
    "avec detection automatique des doublons et rapport d'import detaille."
))
figure_caption(doc, "Figure 27: Gestion des affilies")

heading2(doc, "Formulaire de depot de liquidation (vue affilie)")
body_para(doc, (
    "Le formulaire de depot de liquidation guide l'affilie dans la soumission de "
    "sa demande de retraite. Un systeme d'upload de documents (CIN, attestation, "
    "RIB, acte de naissance) avec validation de taille (max 5 Mo) et barre de "
    "progression permet un depot numerique complet. Une etape de confirmation "
    "recapitule les documents telecharges avant la soumission finale. Un avertissement "
    "clair rappelle l'obligation de se presenter en agence avec les originaux."
))
figure_caption(doc, "Figure 28: Formulaire de depot liquidation")

heading2(doc, "Suivi de dossier (vue affilie)")
body_para(doc, (
    "Une fois la demande deposee, l'affilie accede a une vue de suivi de son dossier "
    "sous forme de timeline interactive. Cinq etapes sont representees : demande "
    "deposee, dossier numerique verifie, verification des pieces originales en agence, "
    "validation finale et liquidation terminee. L'etape en cours est mise en evidence "
    "et un rappel pour les documents physiques est affiche en permanence."
))
figure_caption(doc, "Figure 29: Suivi de dossier (vue affilie)")

heading2(doc, "Gestion des liquidations (vue admin)")
body_para(doc, (
    "L'administrateur dispose d'une vue complete de tous les dossiers de liquidation, "
    "organises en deux onglets : 'Dossiers en attente' et 'Historique'. Des filtres "
    "par statut et une barre de recherche facilitent la navigation. Une modal de "
    "detail permet de consulter les documents joints avec apercu integre. Un autre "
    "modal permet de changer le statut du dossier et d'enregistrer un motif de rejet."
))
figure_caption(doc, "Figure 30: Page de gestion des liquidations (vue admin)")

heading2(doc, "Centre de notifications")
body_para(doc, (
    "Le centre de notifications presente toutes les alertes en temps reel. "
    "Les notifications sont organisees par type (liquidation, paiement, systeme, "
    "securite) et filtrees par statut (toutes, non-lues, lues, corbeille). "
    "Un systeme de corbeille locale (localStorage) permet une suppression douce "
    "avant suppression definitive. Le polling automatique toutes les 30 secondes "
    "et l'affichage d'un toast a chaque nouvelle notification garantissent une "
    "experience utilisateur reactive."
))
figure_caption(doc, "Figure 31: Centre de notifications")

heading2(doc, "Simulation de pension")
body_para(doc, (
    "L'outil de simulation actuarielle permet a l'affilie ou a l'agent de calculer "
    "une estimation de la pension mensuelle en fonction de parametres comme l'age "
    "actuel, l'age de depart souhaite, le salaire mensuel brut et le taux de "
    "cotisation. Le moteur de simulation backend (SimulationEngine) applique les "
    "formules actuarielles de l'Article 6 CIMR."
))
figure_caption(doc, "Figure 32: Simulation de pension")

heading2(doc, "Conclusion")
body_para(doc, (
    "Au terme de ce chapitre, nous avons presente les principales interfaces "
    "graphiques du portail CIMR accompagnees d'explications sur le fonctionnement "
    "de chaque module. L'ensemble des fonctionnalites implementees repond aux "
    "besoins identifies lors de la phase d'analyse et constitue une solution "
    "complete et fonctionnelle pour la gestion des droits a la retraite."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# CONCLUSION ET PERSPECTIVES
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hc = doc.add_heading('Conclusion et perspectives', level=1)
for r in hc.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hc.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_border_bottom(hc._p.getparent() if False else hc, '000000', '6')
doc.add_paragraph()

body_para(doc, (
    "En conclusion, le projet de developpement du Portail de Gestion des Droits a la "
    "Retraite CIMR a ete couronne de succes, repondant de maniere efficace aux besoins "
    "strategiques de modernisation de la Caisse Interprofessionnelle Marocaine de Retraite. "
    "Ce projet a represente un veritable defi technique et organisationnel, mais il a "
    "egalement demontre la capacite de l'equipe a concevoir, developper et mettre en "
    "oeuvre une solution robuste, securisee et evolutive."
))
body_para(doc, (
    "L'architecture microservices adoptee garantit la scalabilite et l'independance "
    "des composants du systeme. La separation des responsabilites entre les neuf "
    "microservices permet des evolutions independantes sans impacter l'ensemble de "
    "la plateforme. La communication asynchrone via Apache Kafka assure la robustesse "
    "du systeme face aux pannes d'infrastructure."
))
body_para(doc, (
    "La securite a ete une preoccupation centrale tout au long du projet, avec la "
    "mise en place de l'authentification JWT, du hachage BCrypt des mots de passe, "
    "du controle d'acces base sur les roles et de la traçabilite via les audit logs. "
    "La conformite RGPD et CNDP a ete integree nativement dans le module d'affiliation."
))
body_para(doc, (
    "L'audit qualite realise sur la codebase a permis d'identifier et de corriger "
    "9 bugs, dont 2 de severite critique impactant directement l'experience utilisateur "
    "(bug de serialisation JSON des notifications et champs manquants dans le modele "
    "de liquidation). La verification TypeScript (tsc --noEmit) confirme l'absence "
    "d'erreurs de compilation dans le frontend."
))

heading2(doc, "Perspectives")
body_para(doc, "Pour les evolutions futures du portail, plusieurs axes d'amelioration sont envisages :")
bullet_item(doc, "Mise en place de migrations de base de donnees automatisees avec Flyway ou Liquibase pour gerer l'evolution du schema en production ;")
bullet_item(doc, "Developpement d'une application mobile (React Native) pour offrir aux affilies un acces depuis leur smartphone ;")
bullet_item(doc, "Integration d'un tableau de bord analytique avance (BI) pour l'analyse des tendances de cotisations et de liquidations ;")
bullet_item(doc, "Mise en place de tests d'integration et de bout en bout automatises pour garantir la non-regression lors des evolutions ;")
bullet_item(doc, "Implementation d'un circuit breaker (Resilience4j) pour une gestion plus robuste des pannes des services tiers ;")
bullet_item(doc, "Extension des capacites de l'agent IA pour la detection automatique de fraudes et l'analyse predictive des dossiers.")

body_para(doc, (
    "En somme, ce projet de fin d'etudes represente un exemple reussi de la maniere "
    "dont les technologies modernes peuvent etre mises au service de la protection "
    "sociale. Il constitue une base solide pour la transformation numerique de la "
    "CIMR et souligne l'importance d'une architecture bien conçue, d'une methodologie "
    "rigoureuse et d'une culture de la qualite dans le developpement logiciel."
))

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# WEBOGRAPHIE
# ════════════════════════════════════════════════════════════════════════
add_logos_header(doc)
hw = doc.add_heading('Webographie', level=1)
for r in hw.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
hw.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_border_bottom(hw, '000000', '6') if False else None
doc.add_paragraph()

refs = [
    ("[1]", "Site officiel CIMR — www.cimr.ma, consulte en 2025"),
    ("[2]", "Spring Boot Documentation — https://spring.io/projects/spring-boot, consulte en 2025"),
    ("[3]", "React Documentation — https://react.dev, consulte en 2025"),
    ("[4]", "Apache Kafka Documentation — https://kafka.apache.org/documentation, consulte en 2025"),
    ("[5]", "TypeScript Handbook — https://www.typescriptlang.org/docs, consulte en 2025"),
    ("[6]", "JSON Web Tokens — https://jwt.io/introduction, consulte en 2025"),
    ("[7]", "Spring Security Reference — https://docs.spring.io/spring-security, consulte en 2025"),
    ("[8]", "Docker Documentation — https://docs.docker.com, consulte en 2025"),
    ("[9]", "Kubernetes Documentation — https://kubernetes.io/docs, consulte en 2025"),
    ("[10]","Zod Documentation — https://zod.dev, consulte en 2025"),
    ("[11]","Framer Motion — https://www.framer.com/motion, consulte en 2025"),
    ("[12]","FastAPI Documentation — https://fastapi.tiangolo.com, consulte en 2025"),
]

for ref_id, ref_text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(ref_id + ' ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
    r2 = p.add_run(ref_text)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'

# ════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ════════════════════════════════════════════════════════════════════════
out = r"c:\Users\iliass\Desktop\projet-fin-d-etude-main\Rapport_PFE_CIMR_Mharrech_Iliass.docx"
doc.save(out)
print(f"Rapport genere : {out}")
